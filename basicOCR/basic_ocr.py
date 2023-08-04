# import regular expresstion to clean more
from PIL import Image

import pytesseract
from docx import Document


text = pytesseract.image_to_string('englishtext.png')

print(text)


document = Document()
document.add_heading('OCR ภาษาไทย', 0)

p = document.add_paragraph(text)
p.add_run('\n\nby Uncle Engineer')

document.add_page_break()
document.save('demo.docx')



# thai_text = pytesseract.image_to_string('thaitext.png', lang='tha')

# raw_text = thai_text.split(' ')

# clean_text1 = []

# for r in raw_text:
#     if r != '':
#         clean_text1.append(r.strip().replace('\n', ''))

# check_text = 'start'

# company = ''

# for t in clean_text1:
#     if 'ชื่อสถานประกอบการ' in check_text:
#         company = t
#         print(f'Found: {company}')
#     check_text = t

